import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Asistente Legal de Guatemala", page_icon="🇬🇹")

# Título de la aplicación
st.title("Asistente Legal de Guatemala")

# Acceder a las claves de API de los secretos de Streamlit
TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

# Inicializar el historial de chat en el estado de la sesión si no existe
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

def buscar_informacion(query):
    url = "https://google.serper.dev/search"
    payload = json.dumps({
        "q": query + " ley Guatemala"
    })
    headers = {
        'X-API-KEY': SERPER_API_KEY,
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()

def generar_respuesta(prompt, contexto, chat_history):
    url = "https://api.together.xyz/inference"
    history_text = "\n".join([f"Human: {q}\nAI: {a}" for q, a in chat_history])
    
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": f"Eres un asistente legal especializado en las leyes de Guatemala. Responde las preguntas basándote en el contexto proporcionado, tu conocimiento general sobre las leyes de Guatemala y el historial de la conversación. Si no tienes suficiente información, indica que no puedes responder con certeza.\n\nHistorial de la conversación:\n{history_text}\n\nContexto: {contexto}\n\nHuman: {prompt}\n\nAI:",
        "max_tokens": 5512,
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["Human:"]
    })
    headers = {
        'Authorization': f'Bearer {TOGETHER_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()['output']['choices'][0]['text'].strip()

def create_docx(chat_history, fuentes):
    doc = Document()
    doc.add_heading('Asistente Legal de Guatemala - Historial de Conversación', 0)

    for pregunta, respuesta in chat_history:
        doc.add_heading('Pregunta', level=1)
        doc.add_paragraph(pregunta)
        doc.add_heading('Respuesta', level=1)
        doc.add_paragraph(respuesta)

    doc.add_heading('Fuentes', level=1)
    for fuente in fuentes:
        doc.add_paragraph(fuente, style='List Bullet')

    doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes oficiales o un abogado para asuntos legales importantes.')

    return doc

# Interfaz de usuario
st.write("Bienvenido al Asistente Legal de Guatemala. Puedes hacerme preguntas sobre las leyes guatemaltecas.")

# Mostrar el historial de chat
for pregunta, respuesta in st.session_state.chat_history:
    st.text_area("Tú:", value=pregunta, height=100, disabled=True)
    st.text_area("Asistente:", value=respuesta, height=200, disabled=True)

pregunta = st.text_input("Ingresa tu pregunta sobre la ley de Guatemala:")

if st.button("Enviar"):
    if pregunta:
        with st.spinner("Buscando información y generando respuesta..."):
            # Buscar información relevante
            resultados_busqueda = buscar_informacion(pregunta)
            contexto = "\n".join([result.get('snippet', '') for result in resultados_busqueda.get('organic', [])])

            # Generar respuesta
            respuesta = generar_respuesta(pregunta, contexto, st.session_state.chat_history)

            # Agregar la nueva interacción al historial
            st.session_state.chat_history.append((pregunta, respuesta))

            # Mostrar la nueva respuesta
            st.text_area("Tú:", value=pregunta, height=100, disabled=True)
            st.text_area("Asistente:", value=respuesta, height=200, disabled=True)

            # Mostrar fuentes
            st.write("Fuentes:")
            fuentes = []
            for resultado in resultados_busqueda.get('organic', [])[:3]:
                fuente = f"{resultado['title']}: {resultado['link']}"
                st.write(f"- [{resultado['title']}]({resultado['link']})")
                fuentes.append(fuente)

            # Crear documento DOCX
            doc = create_docx(st.session_state.chat_history, fuentes)

            # Guardar el documento DOCX en memoria
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)

            # Opción para exportar a DOCX
            st.download_button(
                label="Descargar historial de conversación como DOCX",
                data=docx_file,
                file_name="historial_asistente_legal_guatemala.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    else:
        st.warning("Por favor, ingresa una pregunta.")

# Agregar información en el pie de página
st.markdown("---")
st.markdown("**Nota:** Este asistente utiliza IA para generar respuestas basadas en información disponible en línea. "
            "Siempre verifica la información con fuentes oficiales o un abogado para asuntos legales importantes.")
