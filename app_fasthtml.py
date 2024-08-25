from fasthtml import FastHTML, Param, Response
import requests
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO

app = FastHTML(__name__)

# Configuración de la página
app.set_page_config(page_title="Asistente Legal de Guatemala", page_icon="🇬🇹")

# Título de la aplicación
app.title("Asistente Legal de Guatemala")

# Acceder a las claves de API de los secretos de FastHTML
TOGETHER_API_KEY = app.secrets["TOGETHER_API_KEY"]
SERPER_API_KEY = app.secrets["SERPER_API_KEY"]

@app.route("/")
def index():
    return app.render_template("index.html")

@app.route("/buscar_informacion", methods=["POST"])
def buscar_informacion(query: Param(str)):
    url = "https://google.serper.dev/search"
    payload = json.dumps({
        "q": query + " ley Guatemala"
    })
    headers = {
        'X-API-KEY': SERPER_API_KEY,
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return Response(response.json())

@app.route("/generar_respuesta", methods=["POST"])
def generar_respuesta(prompt: Param(str), contexto: Param(str)):
    url = "https://api.together.xyz/inference"
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": f"Contexto: {contexto}\n\nPregunta: {prompt}\n\nResponde la pregunta basándote en el contexto proporcionado y tu conocimiento general sobre las leyes de Guatemala. Si no tienes suficiente información, indica que no puedes responder con certeza.\n\nRespuesta:",
        "max_tokens": 5512,
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["Pregunta:"]
    })
    headers = {
        'Authorization': f'Bearer {TOGETHER_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return Response(response.json()['output']['choices']['text'].strip())

def create_docx(pregunta, respuesta, fuentes):
    doc = Document()
    doc.add_heading('Asistente Legal de Guatemala', 0)

    doc.add_heading('Pregunta', level=1)
    doc.add_paragraph(pregunta)

    doc.add_heading('Respuesta', level=1)
    doc.add_paragraph(respuesta)

    doc.add_heading('Fuentes', level=1)
    for fuente in fuentes:
        doc.add_paragraph(fuente, style='List Bullet')

    doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes oficiales o un abogado para asuntos legales importantes.')

    return doc

@app.route("/generar_docx", methods=["POST"])
def generar_docx(pregunta: Param(str), respuesta: Param(str), fuentes: Param(list)):
    doc = create_docx(pregunta, respuesta, fuentes)
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return Response(docx_file.getvalue(), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", download_name="respuesta_legal_guatemala.docx")

if __name__ == "__main__":
    app.run()
