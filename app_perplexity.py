import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO, StringIO
import csv

# Configuración de la página
st.set_page_config(page_title="Asistente Legal de Guatemala", page_icon="🇬🇹", layout="wide")

# Título de la aplicación
st.title("Asistente Legal de Guatemala 🇬🇹⚖️")

# Acceder a las claves de API de los secretos de Streamlit
TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
PERPLEXITY_API_KEY = st.secrets["PERPLEXITY_API_KEY"]

def buscar_informacion(query):
    url = "https://api.perplexity.ai/search"
    payload = json.dumps({
        "q": query + " ley Guatemala"
    })
    headers = {
        'Authorization': f'Bearer {PERPLEXITY_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()

def generar_respuesta(prompt, contexto):
    url = "https://api.together.xyz/inference"
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": f"Contexto: {contexto}\n\nPregunta: {prompt}\n\nResponde la pregunta basándote en el contexto proporcionado y tu conocimiento general sobre las leyes de Guatemala. Si no tienes suficiente información, indica que no puedes responder con certeza.\n\nRespuesta:",
        "max_tokens": 5512,
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["Pregunta:"]
    })
    headers = {
        'Authorization': f'Bearer {TOGETHER_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json()['output']['choices'][0]['text'].strip()

def create_docx(pregunta, respuesta, fuentes):
    doc = Document()
    doc.add_heading('Asistente Legal de Guatemala', 0)
    
    doc.add_heading('Pregunta', level=1)
    doc.add_paragraph(pregunta)
    
    doc.add_heading('Respuesta', level=1)
    doc.add_paragraph(respuesta)
    
    doc.add_heading('Fuentes', level=1)
    for fuente in fuentes:
        doc.add_paragraph(fuente, style='List Bullet')
    
    doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes oficiales o un abogado para asuntos legales importantes.')
    
    return doc

def create_csv(pregunta, respuesta):
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Pregunta", "Respuesta"])
    writer.writerow([pregunta, respuesta])
    return output.getvalue().encode('utf-8')


# Interfaz de usuario
pregunta = st.text_input("Ingresa tu pregunta sobre la ley de Guatemala:")

if st.button("Obtener respuesta"):
    if pregunta:
        with st.spinner("Buscando información y generando respuesta..."):
            # Buscar información relevante
            resultados_busqueda = buscar_informacion(pregunta)
            contexto = "\n".join([result.get('snippet', '') for result in resultados_busqueda.get('organic', [])])
            
            # Generar respuesta
            respuesta = generar_respuesta(pregunta, contexto)
            
            # Mostrar respuesta
            st.write("Respuesta:")
            st.write(respuesta)
            
            # Mostrar fuentes
            st.write("Fuentes:")
            fuentes = []
            for resultado in resultados_busqueda.get('organic', [])[:3]:
                fuente = f"{resultado['title']}: {resultado['link']}"
                st.write(f"- [{resultado['title']}]({resultado['link']})")
                fuentes.append(fuente)
            
            # Crear documento DOCX
            doc = create_docx(pregunta, respuesta, fuentes)
            
            # Guardar el documento DOCX en memoria
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            # Crear CSV
            csv_file = create_csv(pregunta, respuesta)
            
            # Opción para exportar a DOCX
            st.download_button(
                label="Descargar resultados como DOCX",
                data=docx_file,
                file_name="respuesta_legal_guatemala.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            
            # Opción para exportar a CSV
            st.download_button(
                label="Descargar resultados como CSV",
                data=csv_file,
                file_name="respuesta_legal_guatemala.csv",
                mime="text/csv",
            )
            
    else:
        st.warning("Por favor, ingresa una pregunta.")

# Agregar información en el pie de página
st.markdown("---")
st.markdown("**Nota:** Este asistente utiliza IA para generar respuestas basadas en información disponible en línea. "
            "Siempre verifica la información con fuentes oficiales o un abogado para asuntos legales importantes.")
